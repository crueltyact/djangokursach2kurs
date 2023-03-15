import datetime
import os.path
import os.path
import random
import shutil
from difflib import SequenceMatcher
from os.path import join
from pathlib import Path

import boto3 as boto3
import numpy as np
import pandas as pd
import requests
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.decorators import login_required
from django.core.files.storage import FileSystemStorage
from django.http import HttpResponseForbidden, FileResponse
from django.shortcuts import render, redirect
from docxtpl import DocxTemplate
from lxml import etree
from transliterate import translit

from docx import Document as Doc

from excel_to_doc_parser.models import CustomUser, Role, Document
from parser_server.settings import BASE_DIR, MEDIA_ROOT

PLANE_PATH = join(MEDIA_ROOT, "excel", "planes", "18048 09.03.01 WEB OFO 2022.xlsx")
MATRIX_PATH = join(MEDIA_ROOT, "excel", "matrices", "Matrix WEB.xlsx")
TEMPLATE_PATH = join(BASE_DIR, "excel_to_doc_parser", "py", "template.docx")


def check_number(num):
    if num % 10 == 1 and num != 11:
        return '1'
    elif 1 < num % 10 < 5 and (num > 19 or num < 5):
        return '2'
    else:
        return '3'


@login_required(login_url='/login/')
def index(request):
    context = {}
    if request.user.is_authenticated:
        context = {"hello": "hello", "custom_user": CustomUser.objects.get(user=request.user)}
        context["role"] = Role.objects.get(pk=context["custom_user"].role_id)
        if context["custom_user"].role_id == 1:
            if request.method == "POST":
                fs = FileSystemStorage()
                folder = MEDIA_ROOT
                for filename in os.listdir(folder):
                    file_path = os.path.join(folder, filename)
                    if filename == ".gitkeep":
                        continue
                    try:
                        if os.path.isfile(file_path) or os.path.islink(file_path):
                            os.unlink(file_path)
                        elif os.path.isdir(file_path):
                            shutil.rmtree(file_path)
                    except Exception as e:
                        print('An error appear ' + str(e))
    else:
        return HttpResponseForbidden()
    return render(request, "main.html", context)


def get_current_disciplines() -> list:
    disciplines = []
    data = parse_plane(PLANE_PATH)["Основные дисциплины"]
    header = get_header(PLANE_PATH)
    for i, row in data[header['Название дисциплины']].items():
        if not ("блок" in row.lower() or "часть" in row.lower() or "дисциплины" in row.lower()):
            if "*" in row:
                row = row[:row.find("*")].strip()
            disciplines.append(row)
    return disciplines


def get_profile_name() -> str:
    df = pd.read_excel(PLANE_PATH, header=None, index_col=None)
    data = df.dropna(axis="columns", how="all").dropna(axis="rows", how="all")
    for _, row in data.iterrows():
        if "Профиль" in np.array2string(row.values):
            for cell in row.values:
                if isinstance(cell, str) and "Профиль:" in cell:
                    return cell.replace("Профиль:", "").strip()
    return ""


def get_program_code() -> str:
    df = pd.read_excel(PLANE_PATH, header=None, index_col=None)
    data = df.dropna(axis="columns", how="all").dropna(axis="rows", how="all")
    for _, row in data.iterrows():
        if "по направлению подготовки" in np.array2string(row.values):
            for cell in row.values:
                if isinstance(cell, str) and "по направлению подготовки" in cell:
                    return cell.replace("по направлению подготовки", "").strip()
    return ""


def parse_matrix(filename) -> (dict, dict):
    frame = pd.read_excel(filename, header=None, index_col=None)
    header = {}
    data = {}
    all_competencies = {}
    header_row = frame.loc[frame[0] == "КОМПЕТЕНЦИИ"]
    header_row.columns = pd.RangeIndex(header_row.columns.size)
    header_row.reset_index(drop=True, inplace=True)
    header_row = header_row.iloc[0].dropna().to_frame().T
    for i, key in enumerate(header_row.values[0], 1):
        if not pd.isna(key):
            header[key] = i
    frame = frame.drop(range(frame.loc[frame[0] == "КОМПЕТЕНЦИИ"].head().index[0] + 1))
    frame = frame.dropna(axis="columns", how="all")
    frame.reset_index(drop=True, inplace=True)
    all_comps = frame[frame.columns[:2]]
    all_comps.columns = ['competency', 'indicator']
    all_comps.reset_index(drop=True, inplace=True)
    universal_comp_end = \
        all_comps.loc[all_comps['competency'] == "Общепрофессиональные компетенции и индикаторы"].head().index[0]
    common_prof_comp_end = \
        all_comps.loc[all_comps['competency'].str.strip() == "Профессиональные компетенции и индикаторы"].head().index[
            0]
    relation_matrix = frame[frame.columns[2:]]
    relation_matrix.drop([0, universal_comp_end, common_prof_comp_end], axis=0, inplace=True)
    relation_matrix.reset_index(drop=True, inplace=True)
    all_comps.drop([0, universal_comp_end, common_prof_comp_end], axis=0, inplace=True)
    all_comps.reset_index(drop=True, inplace=True)
    all_competencies_names = all_comps['competency'].dropna()
    all_competencies_names.reset_index(drop=True, inplace=True)
    all_indicators = all_comps['indicator']
    disciplines = header
    disciplines.pop('КОМПЕТЕНЦИИ', None)
    disciplines.pop("ИНДИКАТОРЫ", None)
    for key, value in disciplines.items():
        relation_column = relation_matrix.iloc[:, value - 3]
        data[key] = {}
        all_competencies[key] = []
        data[key]['universal_competencies'] = []
        data[key]['general_professional_competencies'] = []
        data[key]['professional_competencies'] = []

        current_discipline_relation = pd.concat([all_indicators, relation_column], axis=1)
        current_discipline_relation.columns = ['indicator', 'value']
        current_discipline_relation = current_discipline_relation[current_discipline_relation['value'].notna()]
        all_comp_for_disc = {'competency': [], 'indicator': []}
        for i, values in current_discipline_relation.iterrows():
            if i % 3 == 2:
                comp_name_index = i - 2
            elif i % 3 == 1:
                comp_name_index = i - 1
            else:
                comp_name_index = i
            all_comp_for_disc['competency'].append(all_competencies_names.iloc[comp_name_index // 3])
            all_comp_for_disc['indicator'].append(values['indicator'])
        all_comp_for_disc_df = pd.DataFrame.from_dict(all_comp_for_disc)
        all_comp_for_disc_df = all_comp_for_disc_df.groupby(['competency'])['indicator'].apply(list)
        all_comp_for_disc_df = all_comp_for_disc_df.to_frame()
        for i, values in all_comp_for_disc_df.iterrows():
            indicators = []
            for indicator in values['indicator']:
                indicators.append([
                    indicator.split(" ")[0],
                    ' '.join(word for word in indicator.split(" ")[1:]).strip()
                ])
            all_competencies[key].append({
                'competency_code': i.split(".")[0],
                'competency_name': ' '.join(word for word in i.split(" ")[1:]).strip(),
                'indicators': indicators
            })
            if i.startswith("УК"):
                data[key]['universal_competencies'].append({
                    'competency_code': i.split(".")[0],
                    'competency_name': ' '.join(word for word in i.split(" ")[1:]).strip(),
                    'indicators': indicators
                })
            elif i.startswith("ОПК"):
                data[key]['general_professional_competencies'].append({
                    'competency_code': i.split(".")[0],
                    'competency_name': ' '.join(word for word in i.split(" ")[1:]).strip(),
                    'indicators': indicators
                })
            else:
                data[key]['professional_competencies'].append({
                    'competency_code': i.split(".")[0],
                    'competency_name': ' '.join(word for word in i.split(" ")[1:]).strip(),
                    'indicators': indicators
                })
    return data, all_competencies


def hours_to_zet(z):
    h = round(z / 36, 1)
    if h == int(h):
        return int(h)
    else:
        return h


def number_to_words(n):
    less_than_ten = {1: 'первом', 2: 'втором', 3: 'третьем', 4: 'четвёртом',
                     5: 'пятом', 6: 'шестом', 7: 'седьмом', 8: 'восьмом',
                     9: 'девятом'}
    ten = {10: 'десятом'}
    from_eleven_to_nineteen = {11: 'одиннадцатом', 12: 'двенадцатом',
                               13: 'тринадцатом', 14: 'четырнадцатом',
                               15: 'пятнадцатом', 16: 'шестнадцатом',
                               17: 'семнадцатом', 18: 'восемнадцатом',
                               19: 'девятнадцатом'}
    n1 = n % 10
    n2 = n - n1
    if n < 10:
        return less_than_ten.get(n)
    elif 10 < n < 20:
        return from_eleven_to_nineteen.get(n)
    elif n >= 10 and n in ten:
        return ten.get(n)
    else:
        return ten.get(n2) + ' ' + less_than_ten.get(n1)


@login_required(login_url='/login/')
def documents(request):
    context = {"documents": Document.objects.filter(user_id=request.user.id),
               "custom_user": CustomUser.objects.get(user=request.user), "disciplines": get_current_disciplines()}
    context["role"] = Role.objects.get(pk=context["custom_user"].role_id)
    if request.method == "POST":
        if request.POST.get("generate"):
            folder = join(str(BASE_DIR), "excel_to_doc_parser/media/generated_files/docx")
            data_df = parse_plane(PLANE_PATH)["Основные дисциплины"]
            header = get_header(PLANE_PATH)
            hours = {}
            for i, row in data_df[header['Название дисциплины']].items():
                if not ("блок" in row.lower() or "часть" in row.lower() or "дисциплины" in row.lower()):
                    if "*" in row:
                        row = row[:row.find("*")].strip()
                    hours[row] = {}
                    hours[row]["lections"] = []
                    hours[row]["seminars"] = []
                    hours[row]["labs"] = []
                    hours[row]["srs"] = []
                    hours[row]["exam"] = []
                    hours[row]["test"] = []
                    if not pd.isna(data_df.iloc[i - 1][header["Лекции"]]):
                        hours[row]["lections"].append(data_df.iloc[i - 1][header["Лекции"]])
                    if not pd.isna(data_df.iloc[i - 1][header["Семинары и практические занятия"]]):
                        hours[row]["seminars"].append(data_df.iloc[i - 1][header["Семинары и практические занятия"]])
                    if not pd.isna(data_df.iloc[i - 1][header["Лабораторные работы"]]):
                        hours[row]["labs"].append(data_df.iloc[i - 1][header["Лабораторные работы"]])
                    if not pd.isna(data_df.iloc[i - 1][header["СРС"]]):
                        hours[row]["srs"].append(data_df.iloc[i - 1][header["СРС"]])
                    if not pd.isna(data_df.iloc[i - 1][header["Экзамены"]]):
                        hours[row]["exam"].append(data_df.iloc[i - 1][header["Экзамены"]])
                    if not pd.isna(data_df.iloc[i - 1][header["Зачёты"]]):
                        hours[row]["test"].append(data_df.iloc[i - 1][header["Зачёты"]])
            for filename in os.listdir(folder):
                file_path = os.path.join(folder, filename)
                if filename == ".gitkeep":
                    continue
                try:
                    if os.path.isfile(file_path) or os.path.islink(file_path):
                        os.unlink(file_path)
                    elif os.path.isdir(file_path):
                        shutil.rmtree(file_path)
                except Exception as e:
                    print('An error appear ' + str(e))
            data, all_competencies = parse_matrix(MATRIX_PATH)
            discipline = Document.objects.get(pk=request.POST.get('document')).document_name
            try:
                context_plane_df = data_df[data_df[header['Название дисциплины']] == discipline]
            except KeyError:
                for error_key in data_df[header['Название дисциплины']]:
                    if SequenceMatcher(None, discipline, error_key).ratio() >= 0.75:
                        context_plane_df = data_df[data_df[header['Название дисциплины']] == error_key]
                        break
            context_plane = {
                'intensity_ZET': int(context_plane_df[header['Всего, ЗЕТ']].values[0]),
                'intensity_hours': int(context_plane_df[header['ВСЕГО по структуре']].values[0]),
                'total_homework_hours': int(context_plane_df[header['СРС']].values[0]), 'courses': []
            }
            context_plane['intensity_ZET_check'] = check_number(context_plane['intensity_ZET'])
            context_plane['intensity_hours_check'] = check_number(context_plane['intensity_hours'])
            context_plane['total_homework_hours_check'] = check_number(context_plane['total_homework_hours'])

            for key in header:
                if "семестр" in key:
                    if not pd.isna(context_plane_df[header[key]].values[0]):
                        context_plane['courses'].append({
                            'ZET': hours_to_zet(int(context_plane_df[header[key]].values[0]) + int(
                                context_plane_df[header['СРС']].values[0])),
                            'hours': int(context_plane_df[header[key]].values[0]) + int(
                                context_plane_df[header['СРС']].values[0]),
                            'homework_time': int(context_plane_df[header['СРС']].values[0]),
                            'semester': number_to_words(int(key.split(" ")[0])),
                            'course': number_to_words(int(round(int(key.split(" ")[0]) / 2 + 0.1))),
                            'exam': context_plane_df[header['Экзамены']].values[0] if not pd.isna(
                                context_plane_df[header['Экзамены']].values[0]) and key.split(" ")[0] in context_plane_df[header['Экзамены']].values[0] else "",
                            'test': context_plane_df[header['Зачёты']].values[0] if not pd.isna(
                                context_plane_df[header['Зачёты']].values[0]) and key.split(" ")[0] in context_plane_df[header['Зачёты']].values[0] else ""
                        })
            for i, _ in enumerate(context_plane['courses']):
                context_plane['courses'][i]['ZET_check'] = check_number(context_plane['courses'][i]['ZET'])
                context_plane['courses'][i]['hours_check'] = check_number(context_plane['courses'][i]['hours'])
                context_plane['courses'][i]['homework_time_check'] = check_number(
                    context_plane['courses'][i]['homework_time'])
            try:
                context_plane["hours"] = hours[discipline]
            except KeyError:
                for error_key in data_df[header['Название дисциплины']]:
                    if SequenceMatcher(None, discipline, error_key).ratio() >= 0.75:
                        context_plane["hours"] = hours[error_key]
                        break
            doc = DocxTemplate(TEMPLATE_PATH)
            data = dict(data[discipline], **xml_parser(request))
            data['all_comp'] = all_competencies[discipline]
            data['dean'] = "Д.Г. Демидов"
            data['head_of_faculty'] = "Е.В. Пухова"
            data['rop'] = "М.В. Даньшина"
            data["program_name"] = discipline
            data["current_year"] = datetime.date.today().year
            data["program_code"] = Document.objects.get(pk=request.POST.get('document')).program_code
            data["profile_name"] = Document.objects.get(pk=request.POST.get('document')).profile_name
            data["year_start"] = data['current_year']
            doc.render(dict(data, **context_plane))
            for i in range(len(doc.tables)):
                table = doc.tables[i]._tbl
                for row in doc.tables[i].rows:
                    if len(row.cells) and len(row.cells[0].text.strip()) == 0 and len(set(row.cells)) == 1:
                        table.remove(row._tr)
            doc.save(join(str(BASE_DIR), folder, "{}.docx".format(discipline)))
            context['path'] = join(folder, "{}.docx".format(discipline))
            context['name'] = discipline + '.docx'
            return redirect("/download/?file={}&name=".format(context['path'], context["name"]))
        program_name = request.POST.get("program_name")
        link = request.POST.get("link")
        status = request.POST.get("status")
        user = request.user.id
        new_document = Document(link_to_xml_id=link, status_id=status, user_id=user, document_name=program_name,
                                profile_name=get_profile_name(), program_code=get_program_code())
        new_document.save()
        return redirect('/documents')
    return render(request, "./docx_creation/document.html", context)


def main_info_parser(part, content):
    data = {}
    for field in part:
        data[str(field.tag)] = field.text
    content["main_info"] = data
    return content


def files_parser(part, content):
    data = {}
    for field in part:
        data[str(field.tag)] = field.text
    content["files"] = data
    return content


def targets_parser(part, content):
    data = []
    for field in part:
        if not field.text:
            part.remove(field)
    for i, field in enumerate(part):
        data.append(field.text.strip() + (";" if i < len(part) - 1 else ""))
    content["targets"] = data
    return content


def tasks_parser(part, content):
    data = []
    for field in part:
        if not field.text:
            part.remove(field)
    for i, field in enumerate(part):
        data.append(field.text.strip() + (";" if i < len(part) - 1 else ""))
    content["tasks"] = data
    return content


def sections_parser(part, content):
    for i, field in enumerate(part):
        content[str(field[0].text)] = field[1].text
    content["section"] = content
    return content


def disciplines_parser(part, content):
    data = []
    for field in part:
        if not field.text:
            part.remove(field)
    for i, field in enumerate(part):
        if field.text:
            data.append(field.text.strip() + (";" if field != part[-1] else "."))
    content["disciplines"] = data
    return content


def sections_content_parser(part, content):
    data = []
    for i, field in enumerate(part):
        if field[0].text != "#TODO":
            data.append([field[0].text, int(field[1].text), int(field[2].text), int(field[3].text), int(field[4].text),
                         field[5].text])
    content["sections"] = data
    print(data)
    return content


def marks_parser(part, content):
    data = []
    content[part[0].tag] = part[0].text
    content[part[1].tag] = part[1].text
    content[part[2].tag] = part[2].text
    content[part[3].tag] = part[3].text
    # for i, field in enumerate(part[4]):
    #     content["{}{}".format(field.tag, i)] = field.text
    for i, field in enumerate(part[4]):
        if field[0].text:
            data.append([field[0].text, field[1].text])
    content["marks"] = data
    return content


def literature_parser(part, content):
    data = {}
    for j in range(len(part)):
        books = []
        for i, field in enumerate(part[j]):
            if field.text:
                books.append(field.text.strip())
        data[part[j].tag] = books
    content["literature"] = data
    return content


def software_parser(part, content):
    data = []
    for field in part:
        if not field.text:
            part.remove(field)
    for i, field in enumerate(part):
        data.append(field.text.strip())
    content["software"] = data
    return content


def evaluation_tools_parser(part, content):
    for i, field in enumerate(part):
        content["{}{}".format(field.tag, i)] = field.text
    content["evaluation_tool"] = content
    return content


def tasks_for_students_parser(part, content):
    for i, field in enumerate(part):
        content["{}{}".format(field.tag, i)] = field.text
    content["tasks_for_students"] = content
    return content


def education_technologies_parser(part, content):
    data = {}
    education_technologies_in = []
    education_technologies_out = []
    for field in part[0]:
        if not field.text:
            part[0].remove(field)
    for field in part[1]:
        if not field.text:
            part[1].remove(field)
    for i, field in enumerate(part[0]):
        education_technologies_in.append(field.text.strip() + (";" if field != part[-1] else "."))
    data["education_technologies_in"] = education_technologies_in
    for i, field in enumerate(part[1]):
        education_technologies_out.append(field.text.strip() + (";" if field != part[-1] else "."))
    data["education_technologies_out"] = education_technologies_out
    content["education_technologies"] = data
    return content


def mark_criteries_parser(part, content):
    data = []
    for field in part:
        if not field.text:
            part.remove(field)
    for i, field in enumerate(part):
        data.append(field.text.strip() + (";" if field != part[-1] else "."))
    content["mark_criteries"] = data
    return content


def fos_parser(part, content):
    data = {
        'kr': {},
        'course_work': {},
        'exam_questions': {},
        'test_questions': [],
        'example_exam_questions': [],
        'example_exam_task': '',
        'os': []
    }
    data['exam_questions']['theory'] = []
    data['exam_questions']['tasks'] = []
    data['course_work']['themes'] = []
    data['course_work']['contents'] = []
    for field in part:
        if field.getchildren():
            if field.tag == "kr":
                for i, questions in enumerate(field.getchildren()):
                    data['kr'][i] = []
                    for question in questions:
                        data['kr'][i].append(question.text)
            elif field.tag == 'course_work':
                for theme in field.getchildren()[0]:
                    data['course_work']['themes'].append(theme.text)
                for content in field.getchildren()[1]:
                    data['course_work']['contents'].append(content.text)
            elif field.tag == 'exam_questions':
                for question in field.getchildren()[0]:
                    data['exam_questions']['theory'].append(question.text)
                for tasks in field.getchildren()[1]:
                    data['exam_questions']['tasks'].append(tasks.text)
            elif field.tag == 'os_list':
                for os_list in field.getchildren():
                    data['os'].append([child.text for child in os_list.getchildren()])
            else:
                for question in field:
                    data['test_questions'].append(question.text)
        if field.tag == 'all_os':
            data['all_os'] = field.text
    for key in data:
        data[key] = {k: v for k, v in data[key].items() if v} if isinstance(data[key], dict) else data[key]

    if data['exam_questions']:
        data['example_exam_questions'] = random.sample(data['exam_questions']['theory'], 2)
        data['example_exam_task'] = random.sample(data['exam_questions']['tasks'], 1)
    print(data)
    content['fos'] = data
    return content


def xml_parser(request) -> dict:
    root = etree.fromstring(
        requests.get(download_xml_from_s3(
            request,
            translit("{}.xml".format(
                Document.objects.get(
                    pk=request.POST.get("document")
                ).document_name).replace(" ", "_"), "ru", reversed=True))
        ).content
    )
    content = {}
    functions = {
        "main_info": main_info_parser,
        "files": files_parser,
        "targets": targets_parser,
        "tasks": tasks_parser,
        "sections": sections_parser,
        "disciplines": disciplines_parser,
        "sections_content": sections_content_parser,
        "marks": marks_parser,
        "literature": literature_parser,
        "software": software_parser,
        "evaluation_tools": evaluation_tools_parser,
        "tasks_for_students": tasks_for_students_parser,
        "education_technologies": education_technologies_parser,
        "mark_criteries": mark_criteries_parser,
        "fos": fos_parser
    }
    for part in root:
        content = functions.get(part.tag, lambda: "Invalid tag")(part, content)
    return content


@login_required(login_url='/login/')
def themes(request):
    context = {}
    if request.user.is_authenticated:
        context["custom_user"] = CustomUser.objects.get(user=request.user)
        context["role"] = Role.objects.get(pk=context["custom_user"].role_id)
        if request.method == "GET":
            context["document"] = request.GET.get("document")
        if request.method == "POST":
            context["document"] = request.POST.get("document")
        context["discipline"] = Document.objects.get(pk=context["document"]).document_name
        context["profile"] = Document.objects.get(pk=context["document"]).profile_name
        context["status"] = Document.objects.get(pk=context["document"]).status.status
    return render(request, "./docx_creation/theme.html", context)


@login_required(login_url='/login/')
def document_information(request):
    context = {}
    if request.user.is_authenticated:
        context["custom_user"] = CustomUser.objects.get(user=request.user)
        context["role"] = Role.objects.get(pk=context["custom_user"].role_id)
        context["all_themes"] = get_current_disciplines()
        context["predefined_techs_in_class"] = {"default": ["выполнение лабораторных работ в лабораториях вуза",
                                                            "индивидуальные и групповые консультации студентов преподавателем, в том числе в виде защиты выполненных заданий в рамках самостоятельной работы"],
                                                "optional": [
                                                    "посещение профильных конференций и работа на мастер-классах экспертов и специалистов индустрии"]}
        context["predefined_techs_out_class"] = {
            "default": ["подготовки к выполнению и подготовки к защите лабораторных работ",
                        "подготовки к текущей аттестации",
                        "подготовки к промежуточной аттестации"],
            "optional": ["чтения литературы и освоения дополнительного материала в рамках тематики дисциплины"]}
        context["predefined_criteries_in_methods_for_students"] = {
            "default": ["уровень освоения студентом учебного материала",
                        "умения студента использовать теоретические знания при выполнении практических задач",
                        "сформированность компетенций",
                        "оформление материала в соответствии с требованиями"],
            "optional": []}
        document = Doc(join(BASE_DIR, "excel_to_doc_parser", "py", "Primerny_perechen_otsenochnykh_sredstv.docx"))
        table = document.tables[0]
        context['os_list'] = []
        for row in table.rows[1:]:
            context['os_list'].append({
                'name': row.cells[0].text,
                'description': row.cells[1].text,
                'short_name': row.cells[2].text
            })
        data = parse_plane(PLANE_PATH)["Основные дисциплины"]
        header = get_header(PLANE_PATH)
        hours = {}
        for i, row in data[header['Название дисциплины']].items():
            if not ("блок" in row.lower() or "часть" in row.lower() or "дисциплины" in row.lower()):
                if "*" in row:
                    row = row[:row.find("*")].strip()
                hours[row] = {}
                hours[row]["lections"] = []
                hours[row]["seminars"] = []
                hours[row]["labs"] = []
                hours[row]["srs"] = []
                hours[row]["exam"] = []
                hours[row]["course_work"] = []
                if not pd.isna(data.iloc[i - 1][header["Лекции"]]):
                    hours[row]["lections"].append(data.iloc[i - 1][header["Лекции"]])
                if not pd.isna(data.iloc[i - 1][header["Семинары и практические занятия"]]):
                    hours[row]["seminars"].append(data.iloc[i - 1][header["Семинары и практические занятия"]])
                if not pd.isna(data.iloc[i - 1][header["Лабораторные работы"]]):
                    hours[row]["labs"].append(data.iloc[i - 1][header["Лабораторные работы"]])
                if not pd.isna(data.iloc[i - 1][header["СРС"]]):
                    hours[row]["srs"].append(data.iloc[i - 1][header["СРС"]])
                if not pd.isna(data.iloc[i - 1][header["Экзамены"]]):
                    hours[row]["exam"].append(data.iloc[i - 1][header["Экзамены"]])
                if not pd.isna(data.iloc[i - 1][header["Курсовые работы"]]):
                    hours[row]["course_work"].append(data.iloc[i - 1][header["Курсовые работы"]])
                if not pd.isna(data.iloc[i - 1][header["Курсовые проекты"]]):
                    hours[row]["course_work"].append(data.iloc[i - 1][header["Курсовые проекты"]])
        if request.method == "GET":
            context["document"] = request.GET.get("document")
            context["theme"] = Document.objects.get(pk=request.GET.get("document")).document_name
        if request.method == "POST":
            # context["last_values"] = xml_parser(request)
            context["document"] = request.POST.get("document")
            context["theme"] = Document.objects.get(pk=request.POST.get("document")).document_name
        context["hours"] = hours[context["theme"]]
    return render(request, "./docx_creation/targets.html", context)


def generate_xml(request):
    root = etree.Element("root")
    tree = etree.ElementTree(root)
    main_info = etree.Element("main_info")
    desc = etree.Element("discipline")
    desc.text = Document.objects.get(pk=request.POST.get("document")).document_name
    main_info.append(desc)
    prof = etree.Element("profile")
    prof.text = Document.objects.get(pk=request.POST.get("document")).profile_name
    main_info.append(prof)
    course = etree.Element("course")
    course.text = "#TODO"
    main_info.append(course)
    status = etree.Element("status")
    status.text = Document.objects.get(pk=request.POST.get("document")).status.status
    main_info.append(status)
    elective = etree.Element("elective")
    elective.text = "#TODO"
    main_info.append(elective)
    root.append(main_info)
    files = etree.Element("files")
    files_list = ["rpd", "annotation", "fos", "method", "review", "plan", "matrix", "program"]
    try:
        for element in files_list:
            file = etree.Element(element)
            file.text = "#TODO"
            files.append(file)
    except Exception as exc:
        print(exc)
    root.append(files)
    targets = etree.Element("targets")
    try:
        for element in request.POST.get("targets").split(";"):
            target = etree.Element("target")
            target.text = element
            targets.append(target)
    except Exception as exc:
        print(exc)
    root.append(targets)
    tasks = etree.Element("tasks")
    try:
        for element in request.POST.get("tasks").split(";"):
            task = etree.Element("task")
            task.text = element
            tasks.append(task)
    except Exception as exc:
        print(exc)
    root.append(tasks)
    sections = etree.Element("sections")
    try:
        for element in request.POST.get("all_sections").split(";"):
            if element:
                data = element.split(":")
                print(data)
                section = etree.Element("section")
                section_name = etree.Element("section_name")
                section_name.text = data[0]
                section.append(section_name)
                hours_lections = etree.Element("hours_lections")
                hours_lections.text = data[1]
                section.append(hours_lections)
                hours_labs = etree.Element("hours_labs")
                hours_labs.text = data[2]
                section.append(hours_labs)
                hours_seminars = etree.Element("hours_seminars")
                hours_seminars.text = data[3]
                section.append(hours_seminars)
                hours_srs = etree.Element("hours_srs")
                hours_srs.text = data[4]
                section.append(hours_srs)
                sections.append(section)
    except Exception as exc:
        section = etree.Element("section")
        section_name = etree.Element("section_name")
        section_name.text = "#TODO"
        section.append(section_name)
        hours = etree.Element("hours")
        hours.text = "#TODO"
        section.append(hours)
        sections.append(section)
        raise exc
    root.append(sections)
    disciplines = etree.Element("disciplines")
    try:
        for element in request.POST.get("all_modules").split(";"):
            discipline = etree.Element("discipline")
            discipline.text = element
            disciplines.append(discipline)
    except Exception as exc:
        discipline = etree.Element("discipline")
        discipline.text = "#TODO"
        disciplines.append(discipline)
        print(exc)
    root.append(disciplines)
    sections_content = etree.Element("sections_content")
    try:
        for element in request.POST.get("all_sections").split(";"):
            data = element.split(":")
            section_content = etree.Element("section_content")
            theme = etree.Element("theme")
            theme.text = data[0]
            section_content.append(theme)
            hours_lections = etree.Element("hours_lections")
            hours_lections.text = data[1]
            section_content.append(hours_lections)
            hours_labs = etree.Element("hours_labs")
            hours_labs.text = data[2]
            section_content.append(hours_labs)
            hours_seminars = etree.Element("hours_seminars")
            hours_seminars.text = data[3]
            section_content.append(hours_seminars)
            hours_srs = etree.Element("hours_srs")
            hours_srs.text = data[4]
            section_content.append(hours_srs)
            content = etree.Element("content")
            content.text = data[-1]
            section_content.append(content)
            sections_content.append(section_content)
    except Exception as exc:
        section_content = etree.Element("section_content")
        theme = etree.Element("theme")
        theme.text = "#TODO"
        section_content.append(theme)
        hours = etree.Element("hours")
        hours.text = "#TODO"
        section_content.append(hours)
        content = etree.Element("content")
        content.text = "#TODO"
        section_content.append(content)
        sections_content.append(section_content)
        print(exc)
    root.append(sections_content)
    marks = etree.Element("marks")
    competency = etree.Element("competency")
    competency.text = request.POST.get("competentions")
    marks.append(competency)
    attestation = etree.Element("attestation")
    attestation.text = request.POST.get("attestation")
    marks.append(attestation)
    brs = etree.Element("brs")
    brs.text = request.POST.get("score_system")
    marks.append(brs)
    brs_description = etree.Element("brs_description")
    brs_description.text = request.POST.get("score_system_desc")
    marks.append(brs_description)
    # competencies = etree.Element("competencies")
    # try:
    #     for element in request.POST.get("competencies").split(";") or "":
    #         competency = etree.Element("theme")
    #         competency.text = "#TODO"
    #         competencies.append(competency)
    # except Exception as exc:
    #     competency = etree.Element("theme")
    #     competency.text = "#TODO"
    #     competencies.append(competency)
    #     print(exc)
    # marks.append(competencies)
    intermediate = etree.Element("intermediate")
    try:
        for element in request.POST.get("all_marks").split(";"):
            data = element.split(":")
            mark = etree.Element("mark")
            value = etree.Element("value")
            value.text = data[0]
            mark.append(value)
            characteristics = etree.Element("characteristics")
            characteristics.text = data[1]
            mark.append(characteristics)
            intermediate.append(mark)
    except Exception as exc:
        mark = etree.Element("mark")
        value = etree.Element("value")
        competency.text = "#TODO"
        mark.append(value)
        characteristics = etree.Element("characteristics")
        characteristics.text = "#TODO"
        mark.append(characteristics)
        intermediate.append(mark)
        print(exc)
    marks.append(intermediate)
    root.append(marks)
    literature = etree.Element("literature")
    main = etree.Element("main")
    try:
        for element in request.POST.get("main_lit").split(";"):
            book = etree.Element("book")
            book.text = element
            main.append(book)
    except Exception as exc:
        book = etree.Element("book")
        book.text = "#TODO"
        main.append(book)
        print(exc)
    literature.append(main)
    additional = etree.Element("additional")
    try:
        for element in request.POST.get("extra_lit").split(";"):
            book = etree.Element("book")
            book.text = element
            additional.append(book)
    except Exception as exc:
        book = etree.Element("book")
        book.text = "#TODO"
        additional.append(book)
        print(exc)
    literature.append(additional)
    digital = etree.Element("digital")
    try:
        for element in request.POST.get("digital_lit").split(";") or "":
            resources = etree.Element("resources")
            resources.text = element
            digital.append(resources)
    except Exception as exc:
        resources = etree.Element("resources")
        resources.text = "#TODO"
        digital.append(resources)
        print(exc)
    literature.append(digital)
    root.append(literature)
    software = etree.Element("software")
    try:
        for element in request.POST.get("software").split(";"):
            program = etree.Element("program")
            program.text = element
            software.append(program)
    except Exception as exc:
        program = etree.Element("program")
        program.text = "#TODO"
        software.append(program)
        print(exc)
    root.append(software)
    evaluation_tools = etree.Element("evaluation_tools")
    try:
        for element in request.POST.get("evaluation_tools").split(";"):
            tool = etree.Element("tool")
            tool.text = element
            evaluation_tools.append(tool)
    except Exception as exc:
        tool = etree.Element("tool")
        tool.text = "#TODO"
        evaluation_tools.append(tool)
        print(exc)
    root.append(evaluation_tools)
    tasks_for_students = etree.Element("tasks_for_students")
    try:
        for element in request.POST.get("tasks_from_file").split(";") or "":
            task = etree.Element("task")
            task.text = element
            tasks_for_students.append(task)
    except Exception as exc:
        task = etree.Element("task")
        task.text = "#TODO"
        tasks_for_students.append(task)
        print(exc)
    root.append(tasks_for_students)
    education_technologies = etree.Element("education_technologies")
    education_technologies_in = etree.Element("education_technologies_in")
    education_technologies_out = etree.Element("education_technologies_out")
    try:
        for element in request.POST.getlist("default_tech_in_class"):
            tech = etree.Element("tech")
            tech.text = element
            education_technologies_in.append(tech)
        for element in request.POST.getlist("optional_tech_in_class"):
            tech = etree.Element("tech")
            tech.text = element
            education_technologies_in.append(tech)
        for element in request.POST.getlist("default_tech_out_class"):
            tech = etree.Element("tech")
            tech.text = element
            education_technologies_out.append(tech)
        for element in request.POST.getlist("optional_tech_out_class"):
            tech = etree.Element("tech")
            tech.text = element
            education_technologies_out.append(tech)
    except Exception as exc:
        print(exc)
    education_technologies.append(education_technologies_in)
    education_technologies.append(education_technologies_out)
    root.append(education_technologies)
    mark_criteries = etree.Element("mark_criteries")
    try:
        for element in request.POST.getlist("default_tech_in_class"):
            tech = etree.Element("tech")
            tech.text = element
            mark_criteries.append(tech)
    except Exception as exc:
        print(exc)
    root.append(mark_criteries)
    fos = etree.Element("fos")
    kr = etree.Element("kr")
    course_work = etree.Element("course_work")
    exam_questions = etree.Element("exam_questions")
    test_questions = etree.Element("test_questions")
    os_list = etree.Element("os_list")
    all_os = etree.Element("all_os")
    try:
        if request.POST.getlist('kr'):
            for element in request.POST.getlist("kr"):
                questions = etree.Element("questions")
                for quest in element.split(";"):
                    question = etree.Element("question")
                    question.text = quest
                    questions.append(question)
                kr.append(questions)

        if request.POST.get("course_work_themes"):
            cw_themes = etree.Element("themes")
            for element in request.POST.get("course_work_themes").split(";"):
                theme = etree.Element("theme")
                theme.text = element
                cw_themes.append(theme)
            course_work.append(cw_themes)

        if request.POST.get("course_work_contents"):
            cw_contents = etree.Element("contents")
            for element in request.POST.get("course_work_contents").split(";"):
                content = etree.Element("content")
                content.text = element
                cw_contents.append(content)
            course_work.append(cw_contents)

        if request.POST.get("exam_questions_theory"):
            eq_theory = etree.Element("theory")
            for element in request.POST.get("exam_questions_theory").split(";"):
                question = etree.Element("question")
                question.text = element
                eq_theory.append(question)
            exam_questions.append(eq_theory)

        if request.POST.get("exam_questions_tasks"):
            eq_tasks = etree.Element("tasks")
            for element in request.POST.get("exam_questions_tasks").split(";"):
                task = etree.Element("task")
                task.text = element
                eq_tasks.append(task)
            exam_questions.append(eq_tasks)

        if request.POST.get("test_questions"):
            for element in request.POST.get("test_questions").split(";"):
                question = etree.Element("question")
                question.text = element
                test_questions.append(question)

        if request.POST.get('os'):
            for element in request.POST.get('os').split(";;"):
                if element:
                    os_element = etree.Element("os")
                    name = etree.Element("name")
                    name.text = element.split("::")[0].strip()
                    os_element.append(name)
                    description = etree.Element("description")
                    description.text = element.split("::")[1].strip()
                    os_element.append(description)
                    os_content = etree.Element("content")
                    os_content.text = element.split("::")[2].strip()
                    os_element.append(os_content)
                    os_list.append(os_element)

        if request.POST.get('all_os'):
            all_os.text = request.POST.get("all_os").strip()[:-1]

    except Exception as exc:
        print(exc)
    fos.append(kr)
    fos.append(course_work)
    fos.append(exam_questions)
    fos.append(test_questions)
    fos.append(os_list)
    fos.append(all_os)
    root.append(fos)
    path_to_save = join(str(BASE_DIR), "excel_to_doc_parser\\media\\generated_files\\xml\\{}".format(request.user.id))
    Path(path_to_save).mkdir(parents=True, exist_ok=True)
    # filename = "{}-{}.xml".format(Document.objects.get(pk=request.POST.get("document")).program_name.program_name,
    #                               datetime.date.today().strftime("%m.%d.%Y"))
    filename = translit(
        "{}.xml".format(Document.objects.get(pk=request.POST.get("document")).document_name).replace(" ", "_"),
        "ru", reversed=True)
    tree.write(join(str(BASE_DIR), path_to_save, filename), encoding="UTF-8", xml_declaration=True, pretty_print=True)
    upload_xml_to_s3(request, filename, path_to_save)
    os.remove(join(str(BASE_DIR), path_to_save, filename))


def upload_xml_to_s3(request, filename, filepath):
    session = boto3.session.Session()
    s3 = session.client(
        service_name='s3',
        endpoint_url='https://storage.yandexcloud.net',
        aws_access_key_id=os.environ.get('S3_ACCESS_KEY'),
        aws_secret_access_key=os.environ.get('S3_SECRET_KEY'),
    )
    with open(join(str(BASE_DIR), filepath, filename), "rb") as xml:
        s3.put_object(Bucket=os.environ.get('BUCKET_NAME'), Key='xml/{}/{}'.format(request.user.id, filename),
                      Body=xml.read().decode("UTF-8"))


def download_xml_from_s3(request, filename):
    session = boto3.session.Session()
    s3 = session.client(
        service_name='s3',
        endpoint_url='https://storage.yandexcloud.net',
        aws_access_key_id=os.environ.get('S3_ACCESS_KEY'),
        aws_secret_access_key=os.environ.get('S3_SECRET_KEY'),
    )
    return s3.generate_presigned_url('get_object', Params={
        'Bucket': os.environ.get('BUCKET_NAME'),
        'Key': 'xml/{}/{}'.format(request.user.id, filename)},
                                     ExpiresIn=60)


@login_required(login_url='/login/')
def result(request):
    context = {}
    if request.user.is_authenticated:
        context["custom_user"] = CustomUser.objects.get(user=request.user)
        context["role"] = Role.objects.get(pk=context["custom_user"].role_id)
        context["document"] = request.POST.get("document")
        if request.method == "POST":
            if request.POST.get("end") == "generate":
                generate_xml(request)
            if request.POST.get("save"):
                pass
    return render(request, "./docx_creation/result.html", context)


def download(request):
    file = join(str(BASE_DIR) + "/", request.GET.get('file'))
    response = FileResponse(open(file, 'rb'), as_attachment=True,
                            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Length'] = os.path.getsize(file)
    return response


def login_view(request):
    if request.user.is_authenticated:
        return redirect("/documents/")
    if request.method == "POST":
        username = request.POST.get('login')
        password = request.POST.get('password')
        user = authenticate(request, username=username, password=password)
        if user is not None:
            login(request, user)
            return redirect("/documents/")
        else:
            print("Error")
    return render(request, "authorization.html")


def logout_view(request):
    logout(request)
    if not request.user.is_authenticated:
        return redirect("/")
    return render(request, "authorization.html")


@login_required(login_url='/login/')
def info(request):
    context = {}
    if request.user.is_authenticated:
        context["custom_user"] = CustomUser.objects.get(user=request.user)
        context["role"] = Role.objects.get(pk=context["custom_user"].role_id)
    return render(request, "feedback.html", context)


def get_disciplines_hours(data, header, keys):
    for key in keys:
        for index, row in data[header[key]].items():
            if not pd.isna(row):
                keys[key].append((
                    data.iloc[index - 1].iloc[header["Название дисциплины"] - 1],
                    data.iloc[index - 1].iloc[header[key] - 1]))
    return keys


def get_sem(data, index):
    context = {}
    for i, key in enumerate(data.iloc[:, index:].values[1], index + 1):
        if not pd.isna(key):
            context[key] = i
        else:
            break
    return context


def get_hours(data, index):
    context = {}
    for i, key in enumerate(data.iloc[:, index:].values[1], index + 1):
        if not pd.isna(key) and "курс" not in key:
            context[key] = i
        else:
            break
    return context


def get_courses(data, index):
    context = {}
    for i, key in enumerate(data.iloc[:, index:].values[2], index + 1):
        if not pd.isna(key) and "курс" not in key:
            context[key] = i
        else:
            break
    return context


def get_header(filename):
    header = {}
    df = pd.read_excel(filename, header=None, index_col=None)
    data = df.dropna(axis="columns", how="all")
    disciplines = data.copy()
    header_row = disciplines[disciplines.loc[disciplines[2] == "Шифр"].head().index[0]:
                             disciplines.loc[disciplines[2] == "Шифр"].head().index[0] + 3].dropna(axis="columns",
                                                                                                   how='all')
    header_row.columns = pd.RangeIndex(header_row.columns.size)
    for i, key in enumerate(header_row.values[0], 1):
        if not pd.isna(key):
            if "распределение по семестрам" in key.lower():
                header = dict(**header, **get_sem(header_row, i - 1))
            elif "часы" in key.lower():
                header = dict(**header, **get_hours(header_row, i - 1))
            elif "распределение по курсам" in key.lower():
                header = dict(**header, **get_courses(header_row, i - 1))
            else:
                header[key] = i
    return header


def parse_plane(filename) -> dict:
    context = {}
    df = pd.read_excel(filename, header=None, index_col=None)
    data = df.dropna(axis="columns", how="all")
    disciplines = data.copy()
    start_index = 0
    for index, column in disciplines.items():
        if "Блок 1. Дисциплины (модули)" in np.array2string(column.values):
            start_index = index
    disciplines = disciplines.drop(
        range(data.loc[data[start_index] == "Блок 1. Дисциплины (модули)"].head().index[0] - 1))
    context["Факультативные дисциплины"] = data[3].iloc[
        range(data.loc[data[2].isin(["№ п/п"])].head().index[0], data.iloc[-1:].head().index[0] + 1)]
    disciplines = disciplines.drop(
        range(data.loc[data[2].isin(["№ п/п"])].head().index[0], data.iloc[-1:].head().index[0] + 1))
    for column in data:
        if data[column].isin(["8 семестр\n6 недель"]).any():
            break
    disciplines = disciplines.iloc[:, range(column - 1)]
    disciplines = disciplines.dropna(axis='columns', how="all").dropna(axis="rows", how="all")
    # context["Факультативные дисциплины"] = context["Факультативные дисциплины"].dropna(axis='columns', how='all').dropna(axiss='rows', how='all')
    disciplines.reset_index(drop=True, inplace=True)
    new_header = disciplines.iloc[0].astype(int)
    disciplines = disciplines[1:]
    disciplines.columns = new_header
    context["Основные дисциплины"] = disciplines
    return context
